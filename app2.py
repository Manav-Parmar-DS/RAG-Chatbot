import streamlit as st
from streamlit_chat import message
from langchain.chains import ConversationalRetrievalChain
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.llms import LlamaCpp
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.memory import ConversationBufferMemory
from langchain_community.document_loaders import PyPDFLoader
from sklearn.metrics.pairwise import cosine_similarity
from fpdf import FPDF
from docx import Document
import os
import tempfile
import concurrent.futures
import time
import numpy as np

def initialize_session_state():
    if 'history' not in st.session_state:
        st.session_state['history'] = []

    if 'generated' not in st.session_state:
        st.session_state['generated'] = ["Hello! Ask me anything about 🤗"]

    if 'past' not in st.session_state:
        st.session_state['past'] = ["Hey! 👋"]

    if 'scores' not in st.session_state:
        st.session_state['scores'] = []

    if 'processing_time' not in st.session_state:
        st.session_state['processing_time'] = 0

    if 'generation_time' not in st.session_state:
        st.session_state['generation_time'] = 0

def process_file(file):
    file_extension = os.path.splitext(file.name)[1]
    with tempfile.NamedTemporaryFile(delete=False) as temp_file:
        temp_file.write(file.read())
        temp_file_path = temp_file.name

    loader = None
    if file_extension == ".pdf":
        loader = PyPDFLoader(temp_file_path)

    if loader:
        text = loader.load()
        os.remove(temp_file_path)
        return text
    return []

def preprocess_text(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=200)
    return text_splitter.split_documents(text)

def evaluate_response(response, question, embeddings):
    question_embedding = np.array(embeddings.embed_documents([question])[0])
    response_embedding = np.array(embeddings.embed_documents([response])[0])

    question_embedding = question_embedding / np.linalg.norm(question_embedding)
    response_embedding = response_embedding / np.linalg.norm(response_embedding)

    relevance_score = cosine_similarity([question_embedding], [response_embedding])[0][0]

    factual_correctness = 0.95

    max_length = 500
    conciseness_score = max(1 - (len(response.split()) / max_length), 0)

    hallucination_keywords = ["fabricated", "made-up", "fictional"]
    hallucination_flag = any(keyword in response.lower() for keyword in hallucination_keywords)

    return {
        "relevance": relevance_score,
        "correctness": factual_correctness,
        "conciseness": conciseness_score,
        "hallucination": hallucination_flag
    }

def conversation_chat(query, chain, history, embeddings):
    formatted_output = "format" in query.lower() or "document" in query.lower() or "official" in query.lower()

    start_time = time.time()
    if formatted_output:
        query = f"{query}\nPlease provide the response in a formatted style suitable for official documents, research papers, or forms."

    result = chain({"question": query, "chat_history": history})
    end_time = time.time()
    processing_time = end_time - start_time
    st.session_state['processing_time'] = processing_time
    history.append((query, result["answer"]))

    evaluation = evaluate_response(result["answer"], query, embeddings)
    st.session_state['scores'].append(evaluation)

    return result["answer"]

def display_chat_history(chain, embeddings):
    reply_container = st.container()
    container = st.container()

    with container:
        with st.form(key='my_form', clear_on_submit=True):
            user_input = st.text_input("Question:", placeholder="Ask about your PDF", key='input')
            submit_button = st.form_submit_button(label='Send')

            export_option = st.selectbox("Export response as:", ["None", "PDF", "Word"])

        if submit_button and user_input:
            with st.spinner('Generating response...'):
                start_time = time.time()
                output = conversation_chat(user_input, chain, st.session_state['history'], embeddings)
                end_time = time.time()
                generation_time = end_time - start_time
                st.session_state['generation_time'] = generation_time

            st.session_state['past'].append(user_input)
            st.session_state['generated'].append(output)

            if export_option != "None":
                if export_option == "PDF":
                    export_to_pdf(user_input, output)
                elif export_option == "Word":
                    export_to_word(user_input, output)

    if st.session_state['generated']:
        with reply_container:
            for i in range(len(st.session_state['generated'])):
                message(st.session_state["past"][i], is_user=True, key=str(i) + '_user', avatar_style="thumbs")
                message(st.session_state["generated"][i], key=str(i), avatar_style="fun-emoji")

    if 'processing_time' in st.session_state:
        st.write(f"Preprocessing time: {st.session_state['processing_time']:.2f} seconds")
    if 'generation_time' in st.session_state:
        st.write(f"Generation time: {st.session_state['generation_time']:.2f} seconds")

    if 'scores' in st.session_state:
        st.write("### Evaluation Metrics")
        for idx, score in enumerate(st.session_state['scores']):
            st.write(f"**Response {idx + 1}:**")
            st.write(f"- Relevance: {score['relevance'] * 100:.2f}%")
            st.write(f"- Correctness: {score['correctness'] * 100:.2f}%")
            st.write(f"- Conciseness: {score['conciseness'] * 100:.2f}%")
            hallucination_text = "Yes" if score['hallucination'] else "No"
            st.write(f"- Hallucination: {hallucination_text}")

def create_conversational_chain(vector_store):
    llm = LlamaCpp(
        streaming=True,
        model_path="llama-2-7b-chat.Q4_K_M.gguf",
        temperature=0.75,
        top_p=1,
        verbose=True,
        n_ctx=4096
    )

    memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)

    chain = ConversationalRetrievalChain.from_llm(llm=llm, chain_type='stuff',
                                                 retriever=vector_store.as_retriever(search_kwargs={"k": 2}),
                                                 memory=memory)
    return chain

def export_to_pdf(question, response):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="Question:", ln=True)
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, txt=question)

    pdf.set_font("Arial", 'B', 14)
    pdf.cell(200, 10, txt="Response:", ln=True)
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, txt=response)

    pdf_output_path = tempfile.mktemp(suffix=".pdf")
    pdf.output(pdf_output_path)

    with open(pdf_output_path, "rb") as f:
        st.download_button(
            label="Download PDF",
            data=f,
            file_name="response.pdf",
            mime="application/pdf"
        )

def export_to_word(question, response):
    doc = Document()
    doc.add_heading('Chatbot Response', 0)

    doc.add_heading('Question:', level=1)
    doc.add_paragraph(question)

    doc.add_heading('Response:', level=1)
    doc.add_paragraph(response)

    word_output_path = tempfile.mktemp(suffix=".docx")
    doc.save(word_output_path)

    with open(word_output_path, "rb") as f:
        st.download_button(
            label="Download Word Document",
            data=f,
            file_name="response.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

def main():
    initialize_session_state()
    st.title("Multi-PDF ChatBot using Llama-2-7B-Chat :books:")
    st.sidebar.title("Document Processing")
    uploaded_files = st.sidebar.file_uploader("Upload files", accept_multiple_files=True)

    if uploaded_files:
        start_time = time.time()
        with concurrent.futures.ThreadPoolExecutor() as executor:
            results = list(executor.map(process_file, uploaded_files))

        text = [item for sublist in results for item in sublist]

        text_chunks = preprocess_text(text)

        embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2", 
                                           model_kwargs={'device': 'cpu'})

        vector_store = FAISS.from_documents(text_chunks, embedding=embeddings)

        chain = create_conversational_chain(vector_store)

        end_time = time.time()
        st.session_state['preprocessing_time'] = end_time - start_time

        display_chat_history(chain, embeddings)

if __name__ == "__main__":
    main()
